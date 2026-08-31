"""Fill the Hudl/Wyscout teamsheet template from parsed referee-report data.

The template's tables have a fixed shape, so cells are addressed by index:

    tables[0]  r0: [Competition][NAME]  [Date][DD/MM/YY]
               r1: [Scoreline][X-X]
    tables[1]  r0: [ ][HOMETEAM][ ][AWAYTEAM][ ]
               r1: [No.][Player name][ ][Player name][No.]
               r2..r12  starting XI      (11 rows)
               r13..r21 substitutes      (9 rows)

Column 2 of tables[1] holds the vertically merged "STARTING XI" / "SUBSTITUTES"
labels and is never written to.

The output is never translated: it reproduces the template's own (English) labels
plus the data extracted from the PDF, and nothing else. The app's language setting
affects the website only -- no UI or localized string may ever be written into the
teamsheet, so this module takes no language argument by design.
"""

import copy
import io

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

# The template marks the competition, date, scoreline and team-name placeholders red
# (FF0000). The template must not be edited, so every run written is forced black.
BLACK = RGBColor(0, 0, 0)

STARTERS_START, STARTERS_ROWS = 2, 11
SUBS_START, SUBS_ROWS = 13, 9

HOME_NUMBER, HOME_NAME, AWAY_NAME, AWAY_NUMBER = 0, 1, 3, 4


def set_cell_text(cell, text):
    """Replace a cell's text while keeping its paragraph and font formatting.

    Template placeholders are split across several runs ("P" + "layer" + "n" + "ame"),
    so the first run is reused and the rest removed -- assigning `cell.text` would
    throw away the cell's alignment and font.

    The reused run may carry the placeholder's red colour, so it is forced to black;
    everything else (font, size, alignment) is left untouched.
    """
    paragraph = cell.paragraphs[0]
    if not paragraph.runs:
        paragraph.add_run("")

    run = paragraph.runs[0]
    run.text = str(text)
    run.font.color.rgb = BLACK

    # The paragraph mark carries its own colour; recolour it too, so text typed into
    # the cell afterwards does not come out red.
    mark_properties = paragraph._p.get_or_add_pPr().find(qn("w:rPr"))
    if mark_properties is not None:
        for colour in mark_properties.findall(qn("w:color")):
            colour.set(qn("w:val"), "000000")

    for leftover in paragraph.runs[1:]:
        leftover._element.getparent().remove(leftover._element)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def _clone_row_after(table, index):
    """Duplicate a row, keeping borders and shading, and insert it below."""
    tr = table.rows[index]._tr
    tr.addnext(copy.deepcopy(tr))


def _fill_block(table, start, rows, home_players, away_players):
    """Write one squad block, growing it if a squad is longer than the template."""
    needed = max(len(home_players), len(away_players), rows)
    for offset in range(rows, needed):
        _clone_row_after(table, start + offset - 1)

    for index in range(needed):
        cells = table.rows[start + index].cells
        home = home_players[index] if index < len(home_players) else None
        away = away_players[index] if index < len(away_players) else None
        set_cell_text(cells[HOME_NUMBER], home[0] if home else "")
        set_cell_text(cells[HOME_NAME], home[1] if home else "")
        set_cell_text(cells[AWAY_NAME], away[1] if away else "")
        set_cell_text(cells[AWAY_NUMBER], away[0] if away else "")


def fill(data, template_path):
    """Return the filled teamsheet as .docx bytes."""
    document = Document(template_path)
    header, squads = document.tables[0], document.tables[1]

    set_cell_text(header.rows[0].cells[1], data["competition"])
    set_cell_text(header.rows[0].cells[3], data["date"])
    set_cell_text(header.rows[1].cells[1], data["score"])

    set_cell_text(squads.rows[0].cells[HOME_NAME], data["home"]["name"])
    set_cell_text(squads.rows[0].cells[AWAY_NAME], data["away"]["name"])

    # Substitutes first: growing the starters block would shift these row indices.
    _fill_block(squads, SUBS_START, SUBS_ROWS, data["home"]["subs"], data["away"]["subs"])
    _fill_block(
        squads, STARTERS_START, STARTERS_ROWS, data["home"]["starters"], data["away"]["starters"]
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
